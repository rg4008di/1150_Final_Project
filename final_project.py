"""Random Taco Recipe Book Final Project Sunny Reum. I will be creating a program that will randomly generate taco recipes
 and a picture of a taco."""

# bring in the things
import requests #importing requests module
import docx #importing docx module
from PIL import Image #importing Image from pillow
from io import BytesIO #importing bytesIO from io

url = 'https://taco-1150.herokuapp.com/random/?full_taco=true' # Here i am setting the api in a variable. #url for taco recipe API

pic_url = 'https://api.unsplash.com/photos/random/' # Url for unsplash API

query = 'tacos' # search term for Unsplash API

client_id = 'e73fdc2f4eff1bbeb0ec4505b7cc355b6572b434a09f410702383be7536129aa' # Access Key for unsplash API authorization

query_params = {'client_id': client_id, 'query': query} # API parameters for authorization and search 

response = requests.get(pic_url, params=query_params).json() # putting API response into response variable

full_url = response['urls']['full'] # pulling out full URL from the json file received by call to API

img_data = requests.get(full_url) # putting the full photo into img_data variable

img = BytesIO(img_data.content) # CONVERTING PICTURE TO BYTES AND STORING INFO INTO IMG VARIABLE.


foo = ['First', 'Second', 'Third'] # putting an order to recipe titles by using foo to hold 1st, 2nd, 3rd.

document = docx.Document()  # create a new blank document

document.add_paragraph('Random Taco Cookbook', 'Title') # creating title of document
document.add_picture(img, width=docx.shared.Inches(6), height=docx.shared.Inches(6)) #adding picture from unsplash API.
# brian, i got the random image to work! i didnt have time to work out having the rest of the info change.

document.add_paragraph('Code by: Sunny Reum', 'List Bullet') # adding my name to the document
document.add_paragraph('Image by: TJ Dragotta on Unsplash', 'List Bullet') # adding the image author to the document
document.add_paragraph('Recipes from: https://taco-1150.herokuapp.com/random/?full_taco=true', 'List Bullet') # adding the
# website i took the recipes from.
document.add_page_break() #creating a page break so recipes start on a new page


for i in range(3): # using i for a variable in a 'for loop' to hold count and using range to tell the 'for loop' to run 3 times
    taco_recipe = requests.get(url).json()  # Then creating a variable to request the taco recipe.
    seasoning_name = taco_recipe['seasoning']['name'] # getting the seasoning name from the taco recipe json file
    seasoning = taco_recipe['seasoning']['recipe'] # getting the seasoning Recipe from the taco recipe json file
    condiment_name = taco_recipe['condiment']['name'] # getting the condiment name from the taco recipe json file
    condiment = taco_recipe['condiment']['recipe'] # getting the condiment recipe from the taco recipe json file
    mixin_name = taco_recipe['mixin']['name'] # getting the mixin name from the taco recipe json file
    mixin = taco_recipe['mixin']['recipe'] # getting the mixin recipe from the taco recipe json file
    base_layer_name = taco_recipe['base_layer']['name'] # getting the base layer name from the taco recipe json file
    base_layer = taco_recipe['base_layer']['recipe'] # getting the base layer recipe from the taco recipe json file
    shell_name = taco_recipe['shell']['name'] # getting the taco shell name from the taco recipe json file
    shell = taco_recipe['shell']['recipe'] # getting the taco shell recipe from the taco recipe json file
    document.add_paragraph(f'{foo[i]} Taco Recipe', 'Title') # naming the order of the recipe with the title format
    document.add_paragraph(seasoning_name, 'Heading 1') # adding the seasoning title with the heading 1 format
    document.add_paragraph(seasoning) # adding seasoning recipe text
    document.add_paragraph(condiment_name, 'Heading 1') # adding the condiment title with the heading 1 format 
    document.add_paragraph(condiment) # adding the condiment recipe text
    document.add_paragraph(mixin_name, 'Heading 1') # adding the mixin title with the heading 1 format
    document.add_paragraph(mixin) # adding the mixin recipe text
    document.add_paragraph(base_layer_name, 'Heading 1') # adding the base layer title with the heading 1 format
    document.add_paragraph(base_layer) # adding the base layer recipe text
    document.add_paragraph(shell_name, 'Heading 1') # adding the shell title with heading 1 format
    document.add_paragraph(shell) # adding the shell recipe text
    document.add_page_break() # inserting a page break in the document so the next recipe starts on a new page 

document.save('recipe.docx') # saving the document to a word file
